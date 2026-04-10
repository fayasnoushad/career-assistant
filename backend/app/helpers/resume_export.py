"""
Resume Export Helper
Generates PDF and DOCX exports of resume analysis results
"""

from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_pdf(analysis: dict) -> BytesIO:
    """
    Generate a PDF export of the resume analysis

    Args:
        analysis: Dictionary containing the resume analysis data

    Returns:
        BytesIO: PDF file in memory
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter, topMargin=0.5 * inch, bottomMargin=0.5 * inch
    )

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=24,
        textColor=HexColor("#1e40af"),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=16,
        textColor=HexColor("#1e40af"),
        spaceAfter=12,
        spaceBefore=20,
        fontName="Helvetica-Bold",
    )

    body_style = ParagraphStyle(
        "CustomBody", parent=styles["BodyText"], fontSize=11, spaceAfter=6, leading=14
    )

    # Title
    title = Paragraph("Resume Analysis Report", title_style)
    elements.append(title)

    # Document info
    info_data = [
        ["Filename:", analysis.get("filename", "N/A")],
        ["Upload Date:", analysis.get("upload_date", "N/A").split("T")[0]],
    ]

    if analysis.get("target_role"):
        info_data.append(["Target Role:", analysis["target_role"]])
    if analysis.get("experience_level"):
        info_data.append(["Experience Level:", analysis["experience_level"]])

    info_table = Table(info_data, colWidths=[2 * inch, 4.5 * inch])
    info_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#374151")),
                ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#6b7280")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(info_table)
    elements.append(Spacer(1, 0.3 * inch))

    feedback = analysis.get("feedback", {})

    # Overall Score
    score = feedback.get("overall_score", 0)
    score_color = "#16a34a" if score >= 80 else "#eab308" if score >= 60 else "#dc2626"

    score_heading = Paragraph("Overall Score", heading_style)
    elements.append(score_heading)

    score_text = Paragraph(
        f'<font size="36" color="{score_color}"><b>{score}/100</b></font>',
        ParagraphStyle("ScoreStyle", alignment=TA_CENTER, spaceAfter=20),
    )
    elements.append(score_text)

    # Strengths
    elements.append(Paragraph("✓ Strengths", heading_style))
    for strength in feedback.get("strengths", []):
        elements.append(Paragraph(f"• {strength}", body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Weaknesses
    elements.append(Paragraph("✗ Areas for Improvement", heading_style))
    for weakness in feedback.get("weaknesses", []):
        elements.append(Paragraph(f"• {weakness}", body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Skill Gaps
    elements.append(Paragraph("⚠ Skill Gaps", heading_style))
    for gap in feedback.get("skill_gaps", []):
        gap_text = f"<b>{gap['skill']}</b> ({gap['importance']}): {gap['reason']}"
        elements.append(Paragraph(f"• {gap_text}", body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Improvement Suggestions
    elements.append(Paragraph("💡 Improvement Suggestions", heading_style))
    for i, suggestion in enumerate(feedback.get("improvement_suggestions", []), 1):
        elements.append(Paragraph(f"{i}. {suggestion}", body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Recommended Courses
    elements.append(Paragraph("📚 Recommended Courses", heading_style))
    for course in feedback.get("recommended_courses", []):
        elements.append(Paragraph(f"• {course}", body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Formatting Tips
    elements.append(Paragraph("📝 Formatting Tips", heading_style))
    for tip in feedback.get("formatting_tips", []):
        elements.append(Paragraph(f"• {tip}", body_style))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_docx(analysis: dict) -> BytesIO:
    """
    Generate a DOCX export of the resume analysis

    Args:
        analysis: Dictionary containing the resume analysis data

    Returns:
        BytesIO: DOCX file in memory
    """
    buffer = BytesIO()
    doc = Document()

    # Title
    title = doc.add_heading("Resume Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(30, 64, 175)

    # Document info
    doc.add_paragraph()
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = "Light Grid Accent 1"

    info_table.cell(0, 0).text = "Filename:"
    info_table.cell(0, 1).text = analysis.get("filename", "N/A")
    info_table.cell(1, 0).text = "Upload Date:"
    info_table.cell(1, 1).text = analysis.get("upload_date", "N/A").split("T")[0]

    row_idx = 2
    if analysis.get("target_role"):
        row = info_table.add_row()
        row.cells[0].text = "Target Role:"
        row.cells[1].text = analysis["target_role"]
        row_idx += 1

    if analysis.get("experience_level"):
        row = info_table.add_row()
        row.cells[0].text = "Experience Level:"
        row.cells[1].text = analysis["experience_level"]

    # Make first column bold
    for row in info_table.rows:
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    feedback = analysis.get("feedback", {})

    # Overall Score
    score = feedback.get("overall_score", 0)
    score_heading = doc.add_heading("Overall Score", 1)
    score_heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)

    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_para.add_run(f"{score}/100")
    score_run.font.size = Pt(36)
    score_run.bold = True

    if score >= 80:
        score_run.font.color.rgb = RGBColor(22, 163, 74)  # Green
    elif score >= 60:
        score_run.font.color.rgb = RGBColor(234, 179, 8)  # Yellow
    else:
        score_run.font.color.rgb = RGBColor(220, 38, 38)  # Red

    doc.add_paragraph()

    # Strengths
    heading = doc.add_heading("✓ Strengths", 1)
    heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    for strength in feedback.get("strengths", []):
        doc.add_paragraph(strength, style="List Bullet")

    # Weaknesses
    heading = doc.add_heading("✗ Areas for Improvement", 1)
    heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    for weakness in feedback.get("weaknesses", []):
        doc.add_paragraph(weakness, style="List Bullet")

    # Skill Gaps
    heading = doc.add_heading("⚠ Skill Gaps", 1)
    heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    for gap in feedback.get("skill_gaps", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{gap['skill']}").bold = True
        p.add_run(f" ({gap['importance']}): {gap['reason']}")

    # Improvement Suggestions
    heading = doc.add_heading("💡 Improvement Suggestions", 1)
    heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    for suggestion in feedback.get("improvement_suggestions", []):
        doc.add_paragraph(suggestion, style="List Number")

    # Recommended Courses
    heading = doc.add_heading("📚 Recommended Courses", 1)
    heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    for course in feedback.get("recommended_courses", []):
        doc.add_paragraph(course, style="List Bullet")

    # Formatting Tips
    heading = doc.add_heading("📝 Formatting Tips", 1)
    heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    for tip in feedback.get("formatting_tips", []):
        doc.add_paragraph(tip, style="List Bullet")

    # Save to buffer
    doc.save(buffer)
    buffer.seek(0)
    return buffer
